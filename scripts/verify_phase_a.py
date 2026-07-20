"""Phase A 验收验证脚本 — 创建演示资料并逐一验证每条验收标准。

运行方式:
    python scripts/verify_phase_a.py
"""

import hashlib
import io
import os
import sys
from pathlib import Path

# Add project root to path
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


def _build_pdf_with_text(pages_text: list[str]) -> bytes:
    """Build a valid multi-page PDF with visible text using reportlab.

    Each page contains one line of text that pypdf's extract_text() can read.
    """
    from reportlab.pdfgen import canvas as rl_canvas

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=(612, 792))

    for i, text in enumerate(pages_text):
        c.setFont("Helvetica", 14)
        c.drawString(50, 750, text)
        if i < len(pages_text) - 1:
            c.showPage()

    c.save()
    return buf.getvalue()


CHECK_PASS = "[PASS]"
CHECK_FAIL = "[FAIL]"


def create_demo_data(base_dir: Path, project_id: str):
    """Create complete demo data set with all supported formats + edge cases."""
    source_dir = base_dir / "data" / "projects" / project_id / "source"
    source_dir.mkdir(parents=True, exist_ok=True)
    (source_dir / "reports").mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print("Step 1: Creating Demo Data")
    print("=" * 60)

    # 1. Markdown - multi-level headings
    md = source_dir / "reports" / "Q3_weekly_report.md"
    md.write_text(
        "# Q3 Project Weekly Report - Week 28\n\n"
        "## Progress This Week\n\n"
        "### Frontend\n\n"
        "- Finished login page redesign, added dark mode support\n"
        "- Fixed IE11 compatibility bug #4231\n"
        "- First screen load time: 3.2s -> 1.1s\n\n"
        "### Backend\n\n"
        "- API gateway migrated to K8s, 10% canary traffic\n"
        "- DB slow query optimized: order list P99 850ms -> 120ms\n\n"
        "## Risks and Blockers\n\n"
        "1. [HIGH] Payment gateway certificate expires Aug 1, needs urgent renewal\n"
        "2. [MED] MySQL 8.0 upgrade paused due to charset incompatibility\n\n"
        "## Next Week Plan\n\n"
        "- [ ] Certificate renewal + canary verification\n"
        "- [ ] K8s canary expand to 50%\n"
        "- [ ] Start performance benchmark week\n",
        encoding="utf-8",
    )
    print(f"  + MD: {md.relative_to(base_dir)} ({md.stat().st_size} bytes)")

    # 2. TXT - meeting minutes
    txt = source_dir / "meeting_notes.txt"
    txt.write_text(
        "Project Kickoff Meeting Notes\n\n"
        "Date: 2026-07-20\n"
        "Location: Room 3\n"
        "Attendees: Alice, Bob, Carol, Dave\n\n"
        "Summary:\n"
        "Alice presented the project background, confirming Q4 launch target.\n"
        "Bob recommended React 18 + TypeScript for frontend, FastAPI + PostgreSQL for backend.\n"
        "Carol noted K8s resource quotas need coordination with ops team - apply by Friday.\n"
        "Dave reported competitive analysis, suggesting AI writing assistant as differentiator.\n"
        "All members agreed to the proposal.\n\n"
        "Action Items:\n"
        "1. Alice: Complete project charter v2 by Friday\n"
        "2. Bob: Set up dev environment + README by Wednesday\n"
        "3. Carol: Submit K8s resource request by Thursday\n"
        "4. Dave: Track competitor updates, bi-weekly report\n",
        encoding="utf-8",
    )
    print(f"  + TXT: {txt.relative_to(base_dir)} ({txt.stat().st_size} bytes)")

    # 3. PDF - 3 pages with visible text
    pdf = source_dir / "project_plan.pdf"
    pdf_bytes = _build_pdf_with_text([
        "1. Project Plan - Q3 2026 - Executive Summary",
        "2. Project Plan - Implementation Roadmap and Milestones",
        "3. Project Plan - Risk Assessment and Mitigation Strategies",
    ])
    pdf.write_bytes(pdf_bytes)
    print(f"  + PDF: {pdf.relative_to(base_dir)} (3 pages, {pdf.stat().st_size} bytes)")

    # 4. DOCX - multi-level headings
    from docx import Document as DocxDocument
    docx = source_dir / "requirements_spec.docx"
    d = DocxDocument()
    d.add_heading("Requirements Specification", level=1)
    d.add_paragraph("Version: v1.2 | Date: 2026-07-15 | Author: Product Team")

    d.add_heading("Functional Requirements", level=1)

    d.add_heading("User Management", level=2)
    d.add_paragraph("System shall support admin creation, editing, and deactivation of user accounts.")
    d.add_paragraph("User passwords must meet complexity: 8+ chars, upper+lower+digits.")

    d.add_heading("Role-Based Access", level=2)
    d.add_paragraph("System shall support RBAC with three default roles: Admin, Editor, Viewer.")
    d.add_paragraph("Permission changes shall take effect immediately without re-login.")

    d.add_heading("Document Import", level=2)

    d.add_heading("Supported Formats", level=3)
    d.add_paragraph("System shall support import of MD, TXT, PDF, DOCX, XLSX formats.")
    d.add_paragraph("File size limit: 50MB per file; batch limit: 200 files.")

    d.add_heading("Error Handling", level=3)
    d.add_paragraph("Single file parse failure must not interrupt the overall import process.")
    d.add_paragraph("Failed files shall be recorded with filename and failure reason.")

    d.add_heading("Non-Functional Requirements", level=1)

    d.add_heading("Performance", level=2)
    d.add_paragraph("Single file parsing latency shall be under 5 seconds (for files <= 50MB).")
    d.add_paragraph("Concurrent import shall support 4 files parsed in parallel.")

    d.add_heading("Security", level=2)
    d.add_paragraph("All imported files must undergo path safety check to prevent symlink escape.")
    d.add_paragraph("Parsed content must have SHA-256 hash computed for integrity verification.")

    d.save(str(docx))
    print(f"  + DOCX: {docx.relative_to(base_dir)} ({docx.stat().st_size} bytes)")

    # 5. XLSX - 3 sheets
    from openpyxl import Workbook
    xlsx = source_dir / "Q3_budget.xlsx"
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "Personnel"
    ws1.append(["Role", "Headcount", "Monthly(CNY)", "Months", "Subtotal(CNY)"])
    ws1.append(["Frontend Eng", 3, 25000, 3, 225000])
    ws1.append(["Backend Eng", 4, 28000, 3, 336000])
    ws1.append(["Product Mgr", 1, 30000, 3, 90000])
    ws1.append(["QA Engineer", 2, 22000, 3, 132000])
    ws1.append(["Designer", 1, 24000, 3, 72000])

    ws2 = wb.create_sheet("Infrastructure")
    ws2.append(["Item", "Spec", "Unit(CNY/mo)", "Qty", "Months", "Subtotal(CNY)"])
    ws2.append(["K8s Cluster", "16C32G x4 nodes", 8000, 1, 3, 24000])
    ws2.append(["Cloud DB", "MySQL 8.0 4C16G", 3500, 1, 3, 10500])
    ws2.append(["Object Storage", "1TB + CDN", 1200, 1, 3, 3600])
    ws2.append(["Domain + SSL", "example.com", 200, 1, 3, 600])

    ws3 = wb.create_sheet("Misc")
    ws3.append(["Category", "Description", "Amount(CNY)"])
    ws3.append(["Training", "K8s certification x2 people", 8000])
    ws3.append(["Travel", "Beijing trip 1 week", 5000])
    ws3.append(["Team Building", "Quarterly event", 3000])

    wb.save(str(xlsx))
    print(f"  + XLSX: {xlsx.relative_to(base_dir)} (3 sheets, {xlsx.stat().st_size} bytes)")

    # 6. Corrupt file - binary disguised as .txt
    corrupt = source_dir / "corrupt.txt"
    corrupt.write_bytes(b"\x80\x81\x82\x83\x00\xFF\xFE\xFD")
    print(f"  + Corrupt: {corrupt.relative_to(base_dir)} (binary .txt - expected to fail)")

    # 7. Unsupported format - JSON (not in SUPPORTED_EXTENSIONS)
    json_file = source_dir / "config.json"
    json_file.write_text('{"key": "value"}', encoding="utf-8")
    print(f"  + JSON: {json_file.relative_to(base_dir)} (.json - expected to be skipped)")

    # 8. Symlink escape attempt
    try:
        outside = base_dir / "escape_target.txt"
        outside.write_text("DANGER")
        symlink = source_dir / "escape_link.txt"
        if symlink.exists():
            symlink.unlink()
        symlink.symlink_to(outside.resolve())
        print(f"  + Symlink: {symlink.relative_to(base_dir)} -> external (expected: rejected)")
    except (OSError, PermissionError):
        print(f"  ! Symlink escape test skipped (no permission)")

    return source_dir


def verify_import():
    """Run import_project and verify all acceptance criteria."""
    from app.rag.parsers import import_project
    from app.rag.scanner import is_path_safe
    from app.rag.manifest import compute_sha256

    base_dir = PROJECT_ROOT
    project_id = "demo-project"

    print("\n" + "=" * 60)
    print("Step 2: Running import_project()")
    print("=" * 60)

    result = import_project(project_id, base_dir=base_dir)

    print(f"  Project: {result.project_id}")
    print(f"  Total files: {result.total_files}")
    print(f"  Success: {result.success_count}")
    print(f"  Failed: {result.failure_count}")
    print(f"  Skipped: {result.skipped_count}")
    print(f"  Parsed docs: {len(result.parsed)}")

    all_checks = []
    P = CHECK_PASS
    F = CHECK_FAIL

    # ============================================================
    # AC-1: List success/failure files with metadata
    # ============================================================
    print("\n--- AC-1: File Manifest ---")

    success_files = [sf for sf in result.files if sf.parse_status == "success"]
    failed_files = [sf for sf in result.files if sf.parse_status == "failed"]
    skipped_files = [sf for sf in result.files if sf.parse_status == "unsupported"]

    print(f"\n  Success files ({len(success_files)}):")
    for sf in success_files:
        print(f"    [{sf.format:5s}] {sf.relative_path}  SHA256={sf.sha256[:16]}...  {sf.size_bytes:,}B")

    print(f"\n  Failed files ({len(failed_files)}):")
    for sf in failed_files:
        print(f"    [{sf.format:5s}] {sf.relative_path}  ERROR: {sf.error_message}")

    print(f"\n  Skipped files ({len(skipped_files)}):")
    for sf in skipped_files:
        print(f"    [{sf.format:5s}] {sf.relative_path}")

    ok_sha = all(sf.sha256 is not None and len(sf.sha256) == 64 for sf in success_files)
    ok_size = all(sf.size_bytes is not None and sf.size_bytes > 0 for sf in success_files)

    all_checks.append(("AC-1a", "All safely imported files have valid SHA-256 (64 hex chars)", ok_sha))
    all_checks.append(("AC-1b", "All safely imported files have positive file size", ok_size))
    all_checks.append(("AC-1c", ">= 5 files imported successfully", result.success_count >= 5))
    all_checks.append(("AC-1d", ">= 1 file failed (corrupt.txt)", result.failure_count >= 1))

    # ============================================================
    # AC-2: PDF page number citations
    # ============================================================
    print("\n--- AC-2: PDF Page Number Citations ---")
    pdf_docs = [d for d in result.parsed if d.format == "pdf"]
    if pdf_docs:
        doc = pdf_docs[0]
        pages = sorted(set(ch.page_number for ch in doc.chunks if ch.page_number is not None))
        print(f"  Chunks: {len(doc.chunks)}, Pages: {pages}")
        for ch in doc.chunks:
            print(f"    Page {ch.page_number}: {ch.content[:80]}...")
        all_checks.append(("AC-2", f"PDF has {len(doc.chunks)} pages, page_numbers >= 1",
                           len(doc.chunks) >= 3 and all(p >= 1 for p in pages)))
    else:
        all_checks.append(("AC-2", "PDF parsed with page numbers", False))

    # ============================================================
    # AC-3: DOCX heading path citations
    # ============================================================
    print("\n--- AC-3: DOCX Heading Path Citations ---")
    docx_docs = [d for d in result.parsed if d.format == "docx"]
    if docx_docs:
        doc = docx_docs[0]
        headings = [ch for ch in doc.chunks if ch.heading_path]
        multi_level = [h for h in headings if " > " in h.heading_path]
        print(f"  Total chunks: {len(doc.chunks)}, Heading chunks: {len(headings)}")
        print(f"  Multi-level headings: {len(multi_level)}")
        for ch in doc.chunks[:6]:
            hp = ch.heading_path or "(body)"
            print(f"    [{ch.paragraph_index:2d}] {hp}")
            if len(doc.chunks) > 6:
                print(f"    ... ({len(doc.chunks) - 6} more chunks)")
                break
        all_checks.append(("AC-3a", "DOCX has heading_path attached to chunks", len(headings) > 0))
        all_checks.append(("AC-3b", "DOCX has multi-level paths (e.g. A > B > C)", len(multi_level) > 0))
    else:
        all_checks.append(("AC-3", "DOCX parsed with headings", False))

    # ============================================================
    # AC-4: XLSX sheet/header/cell_range citations
    # ============================================================
    print("\n--- AC-4: XLSX Citations (Sheet/Header/Cell Range) ---")
    xlsx_docs = [d for d in result.parsed if d.format == "xlsx"]
    if xlsx_docs:
        doc = xlsx_docs[0]
        sheets = set(ch.sheet_name for ch in doc.chunks)
        has_headers = any(ch.header_columns for ch in doc.chunks)
        has_range = any(ch.cell_range for ch in doc.chunks)
        print(f"  Sheets: {sheets}")
        print(f"  Has headers: {has_headers}, Has cell_ranges: {has_range}")
        for sheet_name in sorted(sheets):
            sheet_chunks = [ch for ch in doc.chunks if ch.sheet_name == sheet_name]
            hdr = sheet_chunks[0].header_columns if sheet_chunks else []
            print(f"  [{sheet_name}] header={hdr}  rows={len(sheet_chunks)}")
        all_checks.append(("AC-4a", "XLSX has 3 sheets", len(sheets) == 3))
        all_checks.append(("AC-4b", "XLSX chunks have header_columns", has_headers))
        all_checks.append(("AC-4c", "XLSX chunks have cell_range", has_range))
    else:
        all_checks.append(("AC-4", "XLSX parsed with sheets/headers/range", False))

    # ============================================================
    # AC-5: MD section title citations
    # ============================================================
    print("\n--- AC-5: Markdown Section Title Citations ---")
    md_docs = [d for d in result.parsed if d.format == "md"]
    if md_docs:
        doc = md_docs[0]
        levels = set(ch.heading_level for ch in doc.chunks if ch.heading_level is not None)
        print(f"  Chunks: {len(doc.chunks)}, Heading levels present: {sorted(levels)}")
        for ch in doc.chunks:
            prefix = f"H{ch.heading_level}" if ch.heading_level else "  -"
            print(f"    {prefix:4s} {ch.section_title or '(preamble)'}  L{ch.line_start}-L{ch.line_end}")
        all_checks.append(("AC-5a", "MD has H1, H2, H3 headings", {1, 2, 3}.issubset(levels)))
        all_checks.append(("AC-5b", "MD chunks have line_start/line_end", all(
            ch.line_start is not None for ch in doc.chunks
        )))
    else:
        all_checks.append(("AC-5", "MD parsed with section titles", False))

    # ============================================================
    # AC-6: TXT paragraph citations
    # ============================================================
    print("\n--- AC-6: TXT Paragraph Line Citations ---")
    txt_docs = [d for d in result.parsed if d.format == "txt" and d.chunks]
    if txt_docs:
        doc = txt_docs[0]
        print(f"  Paragraphs: {len(doc.chunks)}")
        for ch in doc.chunks:
            print(f"    L{ch.line_start:02d}-L{ch.line_end:02d}: {ch.content[:60]}...")
        line_starts = [ch.line_start for ch in doc.chunks]
        increasing = all(
            line_starts[i] < line_starts[i + 1] for i in range(len(line_starts) - 1)
        )
        all_checks.append(("AC-6a", "TXT paragraphs have line ranges", len(doc.chunks) >= 2))
        all_checks.append(("AC-6b", "TXT line numbers are strictly increasing", increasing))
    else:
        all_checks.append(("AC-6", "TXT parsed with paragraph line numbers", False))

    # ============================================================
    # AC-7: Error resilience (corrupt file doesn't crash)
    # ============================================================
    print("\n--- AC-7: Error Resilience ---")
    has_corrupt = any("corrupt" in (sf.relative_path or "").lower() and sf.parse_status == "failed"
                      for sf in result.files)
    good_files_still_ok = result.success_count >= 4
    print(f"  Corrupt file caught: {has_corrupt}")
    print(f"  Good files unaffected (>=4): {good_files_still_ok} (actual={result.success_count})")
    all_checks.append(("AC-7", "Corrupt file fails alone; good files survive", has_corrupt and good_files_still_ok))

    # ============================================================
    # AC-8: SHA-256 integrity check
    # ============================================================
    print("\n--- AC-8: SHA-256 Integrity Check ---")
    sha_ok_cnt = 0
    sha_fail_cnt = 0
    for sf in result.files:
        if sf.parse_status == "success":
            full_path = base_dir / "data" / "projects" / project_id / sf.relative_path
            actual = compute_sha256(full_path)
            if sf.sha256 == actual:
                sha_ok_cnt += 1
            else:
                sha_fail_cnt += 1
                print(f"  {F} SHA mismatch: {sf.relative_path}")
    print(f"  SHA verified: {sha_ok_cnt}, mismatches: {sha_fail_cnt}")
    all_checks.append(("AC-8", f"SHA-256 matches for all {sha_ok_cnt} success files", sha_fail_cnt == 0))

    # ============================================================
    # AC-9: Symlink escape security
    # ============================================================
    print("\n--- AC-9: Symlink Escape Security ---")
    symlink_path = base_dir / f"data/projects/{project_id}/source/escape_link.txt"
    if symlink_path.exists():
        project_dir = (base_dir / f"data/projects/{project_id}").resolve()
        is_safe = is_path_safe(project_dir, symlink_path)
        rejected = next((sf for sf in result.files if "escape_link" in sf.relative_path), None)
        print(f"  is_path_safe: {is_safe}, rejected record: {rejected is not None}")
        all_checks.append((
            "AC-9",
            "Symlink escape is rejected and recorded",
            not is_safe
            and rejected is not None
            and rejected.parse_status == "failed"
            and rejected.error_message == "symlink escape detected",
        ))
    else:
        # Also test the function directly
        outside = base_dir / "escape_target.txt"
        outside.write_text("DANGER")
        project_dir = base_dir / f"data/projects/{project_id}"
        fake_symlink = project_dir / "source" / "fake_escape.txt"
        fake_symlink.write_text("safe_inside")
        is_safe_internal = is_path_safe(project_dir, fake_symlink)

        # Direct defense test: path outside project
        is_safe_external = is_path_safe(project_dir, outside)
        print(f"  Direct test: internal_path safe={is_safe_internal}, external_path safe={is_safe_external}")
        all_checks.append(("AC-9", "External path correctly rejected by is_path_safe",
                           is_safe_internal and not is_safe_external))

    # ============================================================
    # AC-10: Citation reconstruction capability
    # ============================================================
    print("\n--- AC-10: Citation Reconstruction ---")
    citations = []
    for doc in result.parsed:
        fmt = doc.format
        for chunk in doc.chunks:
            cite = ""
            if fmt == "pdf" and chunk.page_number is not None:
                cite = f"{chunk.relative_path}, page {chunk.page_number}"
            elif fmt == "docx" and chunk.heading_path:
                cite = f"{chunk.relative_path}, {chunk.heading_path}"
            elif fmt == "xlsx" and chunk.sheet_name is not None:
                cite = f"{chunk.relative_path}, Sheet '{chunk.sheet_name}', {chunk.cell_range}"
            elif fmt == "md" and chunk.heading_level is not None:
                cite = f"{chunk.relative_path}, {'#' * chunk.heading_level} {chunk.section_title}"
            elif fmt == "txt":
                cite = f"{chunk.relative_path}, lines {chunk.line_start}-{chunk.line_end}"
            else:
                cite = f"{chunk.relative_path}, chunk #{chunk.chunk_index}"
            citations.append(cite)

    print(f"  Total citable chunks: {len(citations)}")
    for c in citations[:8]:
        print(f"    {c}")
    if len(citations) > 8:
        print(f"    ... ({len(citations) - 8} more)")

    all_checks.append(("AC-10", f"{len(citations)} citable chunks across all formats", len(citations) >= 10))

    # ============================================================
    # Summary
    # ============================================================
    print("\n" + "=" * 60)
    print("ACCEPTANCE VERIFICATION SUMMARY")
    print("=" * 60)
    all_pass = True
    for ac_id, desc, passed in all_checks:
        status = P if passed else F
        if not passed:
            all_pass = False
        print(f"  {status}  {ac_id}: {desc}")

    print(f"\n{'=' * 60}")
    if all_pass:
        print(f"{P} ALL ACCEPTANCE CRITERIA PASSED - Phase A is complete.")
    else:
        print(f"{F} SOME CHECKS FAILED - see above.")
    print(f"{'=' * 60}")

    return all_pass


def main():
    base_dir = PROJECT_ROOT
    project_id = "demo-project"

    # Clean up previous demo data
    import shutil
    demo_dir = base_dir / "data" / "projects" / project_id
    if demo_dir.exists():
        shutil.rmtree(demo_dir, ignore_errors=True)
    # Also clean external file from previous symlink test
    escape_file = base_dir / "escape_target.txt"
    if escape_file.exists():
        escape_file.unlink()

    print("Phase A - Document Import & Parsing Verification")
    print(f"Root: {base_dir}")
    print(f"Project: {project_id}\n")

    create_demo_data(base_dir, project_id)
    passed = verify_import()

    sys.exit(0 if passed else 1)


if __name__ == "__main__":
    main()

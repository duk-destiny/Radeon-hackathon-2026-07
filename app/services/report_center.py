"""
Phase H — 报告中心服务

报告草稿创建、编辑、提交审批、版控与导出。
"""

from __future__ import annotations

import sqlite3
import uuid

from io import BytesIO


class ReportCenterService:
    """Manage report drafts, approvals, and exports."""

    def __init__(self, db_path: str):
        self._db_path = db_path

    def _connect(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self._db_path)
        conn.row_factory = sqlite3.Row
        return conn

    # ------------------------------------------------------------------
    # Draft CRUD
    # ------------------------------------------------------------------

    def create_draft(self, project_id: str, title: str, author_id: str, content_md: str = "") -> dict:
        conn = self._connect()
        try:
            rid = str(uuid.uuid4())
            conn.execute(
                "INSERT INTO report_draft (id, project_id, title, content_md, author_id) "
                "VALUES (?, ?, ?, ?, ?)",
                (rid, project_id, title, content_md, author_id),
            )
            conn.commit()
            return self._get_draft(conn, rid)
        finally:
            conn.close()

    def update_draft(self, draft_id: str, title: str | None = None, content_md: str | None = None) -> dict | None:
        conn = self._connect()
        try:
            if title is not None:
                conn.execute(
                    "UPDATE report_draft SET title = ?, updated_at = datetime('now') WHERE id = ?",
                    (title, draft_id),
                )
            if content_md is not None:
                conn.execute(
                    "UPDATE report_draft SET content_md = ?, updated_at = datetime('now') WHERE id = ?",
                    (content_md, draft_id),
                )
            conn.commit()
            return self._get_draft(conn, draft_id)
        finally:
            conn.close()

    def get_draft(self, draft_id: str) -> dict | None:
        conn = self._connect()
        try:
            return self._get_draft(conn, draft_id)
        finally:
            conn.close()

    def list_drafts(self, project_id: str, status: str | None = None) -> list[dict]:
        conn = self._connect()
        try:
            where = "rd.project_id = ?"
            params: list = [project_id]
            if status:
                where += " AND rd.status = ?"
                params.append(status)
            rows = conn.execute(
                f"SELECT rd.*, u.username as author_name "
                f"FROM report_draft rd JOIN user_account u ON rd.author_id = u.id "
                f"WHERE {where} ORDER BY rd.updated_at DESC",
                params,
            ).fetchall()
            return [self._row_to_dict(r) for r in rows]
        finally:
            conn.close()

    def submit_for_approval(self, draft_id: str) -> dict | None:
        conn = self._connect()
        try:
            conn.execute(
                "UPDATE report_draft SET status = 'submitted', updated_at = datetime('now') WHERE id = ?",
                (draft_id,),
            )
            conn.commit()
            return self._get_draft(conn, draft_id)
        finally:
            conn.close()

    def approve_report(
        self, draft_id: str, approver_id: str, decision: str, comment: str = ""
    ) -> dict | None:
        conn = self._connect()
        try:
            new_status = "approved" if decision == "approved" else (
                "rejected" if decision == "rejected" else "draft"
            )
            conn.execute(
                "UPDATE report_draft SET status = ?, updated_at = datetime('now') WHERE id = ?",
                (new_status, draft_id),
            )
            aid = str(uuid.uuid4())
            conn.execute(
                "INSERT INTO report_approval (id, report_id, approver_id, decision, comment) "
                "VALUES (?, ?, ?, ?, ?)",
                (aid, draft_id, approver_id, decision, comment),
            )
            conn.commit()
            return self._get_draft(conn, draft_id)
        finally:
            conn.close()

    def get_approvals(self, draft_id: str) -> list[dict]:
        conn = self._connect()
        try:
            rows = conn.execute(
                "SELECT ra.*, u.username as approver_name, u.display_name "
                "FROM report_approval ra JOIN user_account u ON ra.approver_id = u.id "
                "WHERE ra.report_id = ? ORDER BY ra.created_at DESC",
                (draft_id,),
            ).fetchall()
            return [
                {
                    "id": r["id"],
                    "report_id": r["report_id"],
                    "approver_id": r["approver_id"],
                    "approver_name": r["approver_name"],
                    "decision": r["decision"],
                    "comment": r["comment"],
                    "created_at": r["created_at"],
                }
                for r in rows
            ]
        finally:
            conn.close()

    # ------------------------------------------------------------------
    # Export
    # ------------------------------------------------------------------

    def export_pdf_bytes(self, draft_id: str) -> bytes:
        """Generate PDF bytes for a report draft."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        draft = self.get_draft(draft_id)
        if draft is None:
            raise ValueError("Report draft not found")

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph(f"<b>{draft['title']}</b>", styles["Title"]))
        story.append(Spacer(1, 12))
        for line in draft["content_md"].split("\n"):
            if line.startswith("# "):
                story.append(Paragraph(line[2:], styles["Heading1"]))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["Heading2"]))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:], styles["Heading3"]))
            elif line.strip():
                story.append(Paragraph(line, styles["BodyText"]))
            else:
                story.append(Spacer(1, 6))
        doc.build(story)
        return buf.getvalue()

    def export_docx_bytes(self, draft_id: str) -> bytes:
        """Generate DOCX bytes for a report draft."""
        from docx import Document
        from docx.shared import Pt

        draft = self.get_draft(draft_id)
        if draft is None:
            raise ValueError("Report draft not found")

        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)
        doc.add_heading(draft["title"], level=0)

        for line in draft["content_md"].split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph("")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _get_draft(self, conn: sqlite3.Connection, draft_id: str) -> dict | None:
        row = conn.execute(
            "SELECT rd.*, u.username as author_name "
            "FROM report_draft rd JOIN user_account u ON rd.author_id = u.id "
            "WHERE rd.id = ?",
            (draft_id,),
        ).fetchone()
        return self._row_to_dict(row) if row else None

    @staticmethod
    def _row_to_dict(row: sqlite3.Row) -> dict:
        return {
            "id": row["id"],
            "project_id": row["project_id"],
            "title": row["title"],
            "content_md": row["content_md"],
            "version": row["version"],
            "status": row["status"],
            "author_id": row["author_id"],
            "author_name": row["author_name"] if "author_name" in row.keys() else "",
            "created_at": row["created_at"],
            "updated_at": row["updated_at"],
        }

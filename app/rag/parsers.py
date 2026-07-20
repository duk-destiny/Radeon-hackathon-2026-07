"""Document parsers for Phase A import pipeline.

Supports Markdown, TXT (stdlib), and PDF/DOCX/XLSX (optional deps).
Each parser produces a ParsedDocument with citation-level location chunks.
Parse errors are captured and never propagated — the caller decides
whether to abort or continue.
"""

import re
from pathlib import Path

from app.rag.manifest import (
    ContentChunk,
    ParsedDocument,
    SourceFile,
    ImportResult,
    build_source_file,
)


# ---------------------------------------------------------------------------
# Markdown parser (stdlib only)
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)


def parse_markdown(file_path: Path) -> ParsedDocument:
    """Parse a Markdown file into heading-delimited chunks.

    Text before the first heading is grouped as a chunk with
    ``section_title=None``.  Each heading and its following text
    (up to the next heading of same-or-higher level) forms a chunk.

    Returns:
        ParsedDocument with format ``"md"``.
    """
    text = file_path.read_text(encoding="utf-8")
    lines = text.splitlines(keepends=True)

    # Locate all headings with their positions
    headings: list[tuple[int, int, str]] = []  # (line_index, level, title)
    for m in _HEADING_RE.finditer(text):
        line_no = text[: m.start()].count("\n")
        level = len(m.group(1))
        title = m.group(2).strip()
        headings.append((line_no, level, title))

    chunks: list[ContentChunk] = []

    if not headings:
        # Entire file is one chunk
        body = text.strip()
        if body:
            chunks.append(
                ContentChunk(
                    content=body,
                    relative_path=file_path.name,
                    chunk_index=0,
                    line_start=0,
                    line_end=len(lines),
                    section_title=None,
                    heading_level=None,
                )
            )
        return ParsedDocument(
            relative_path=file_path.name,
            format="md",
            chunks=chunks,
        )

    # Text before first heading
    first_heading_line = headings[0][0]
    if first_heading_line > 0:
        prefix_lines = lines[:first_heading_line]
        prefix_text = "".join(prefix_lines).strip()
        if prefix_text:
            chunks.append(
                ContentChunk(
                    content=prefix_text,
                    relative_path=file_path.name,
                    chunk_index=0,
                    line_start=0,
                    line_end=first_heading_line,
                    section_title=None,
                    heading_level=None,
                )
            )

    # Process each heading as a section
    for i, (h_line, h_level, h_title) in enumerate(headings):
        # Determine end line: next heading of same-or-higher level, or EOF
        end_line = len(lines)
        for j in range(i + 1, len(headings)):
            if headings[j][1] <= h_level:
                end_line = headings[j][0]
                break

        section_text = "".join(lines[h_line:end_line]).strip()
        if section_text:
            chunks.append(
                ContentChunk(
                    content=section_text,
                    relative_path=file_path.name,
                    chunk_index=len(chunks),
                    line_start=h_line,
                    line_end=end_line,
                    section_title=h_title,
                    heading_level=h_level,
                )
            )

    return ParsedDocument(
        relative_path=file_path.name,
        format="md",
        chunks=chunks,
    )


# ---------------------------------------------------------------------------
# Plain-text parser (stdlib only)
# ---------------------------------------------------------------------------

_PARAGRAPH_SPLIT_RE = re.compile(r"\n\s*\n+")


def parse_txt(file_path: Path) -> ParsedDocument:
    """Parse a plain-text file into paragraph-delimited chunks.

    Paragraphs are separated by one or more blank lines.

    Returns:
        ParsedDocument with format ``"txt"``.
    """
    text = file_path.read_text(encoding="utf-8")
    paragraphs = _PARAGRAPH_SPLIT_RE.split(text)

    chunks: list[ContentChunk] = []
    line_offset = 0

    for idx, para in enumerate(paragraphs):
        stripped = para.strip()
        if not stripped:
            line_offset += para.count("\n") + 1
            continue

        para_lines = stripped.count("\n") + 1

        chunks.append(
            ContentChunk(
                content=stripped,
                relative_path=file_path.name,
                chunk_index=idx,
                line_start=line_offset,
                line_end=line_offset + para_lines,
                section_title=None,
                heading_level=None,
            )
        )
        line_offset += para.count("\n") + 2  # paragraph + blank line separator

    return ParsedDocument(
        relative_path=file_path.name,
        format="txt",
        chunks=chunks,
    )


# ---------------------------------------------------------------------------
# PDF parser (optional: pypdf)
# ---------------------------------------------------------------------------


def parse_pdf(file_path: Path) -> ParsedDocument:
    """Parse a PDF file into per-page chunks.

    Requires ``pypdf``.  Each page's extracted text becomes one chunk
    tagged with ``page_number``.

    Returns:
        ParsedDocument with format ``"pdf"``.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError(
            "PDF parsing requires pypdf. Install with: pip install pypdf"
        )

    reader = PdfReader(str(file_path))
    chunks: list[ContentChunk] = []

    for page_idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text and text.strip():
            chunks.append(
                ContentChunk(
                    content=text.strip(),
                    relative_path=file_path.name,
                    chunk_index=page_idx - 1,
                    page_number=page_idx,
                )
            )

    return ParsedDocument(
        relative_path=file_path.name,
        format="pdf",
        chunks=chunks,
    )


# ---------------------------------------------------------------------------
# DOCX parser (optional: python-docx)
# ---------------------------------------------------------------------------


def _is_heading_style(style_name: str) -> bool:
    """Check if a paragraph style name indicates a heading."""
    if not style_name:
        return False
    lowered = style_name.lower()
    return lowered.startswith("heading") or lowered.startswith("head")


def _extract_heading_level(style_name: str) -> int:
    """Try to extract a heading level number from a style name."""
    match = re.search(r"(\d+)", style_name)
    return int(match.group(1)) if match else 1


def parse_docx(file_path: Path) -> ParsedDocument:
    """Parse a DOCX file into paragraph chunks with heading tracking.

    Requires ``python-docx``.  Heading paragraphs provide a
    ``heading_path`` (e.g. ``"1. Introduction > 1.1 Overview"``).
    Non-heading paragraphs are tagged with the current heading path.

    Returns:
        ParsedDocument with format ``"docx"``.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "DOCX parsing requires python-docx. Install with: pip install python-docx"
        )

    doc = Document(str(file_path))
    chunks: list[ContentChunk] = []
    heading_stack: list[tuple[int, str]] = []  # [(level, title), ...]

    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue

        style_name = para.style.name if para.style else ""

        if _is_heading_style(style_name):
            level = _extract_heading_level(style_name)
            # Pop headings of equal or deeper level
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            heading_stack.append((level, text.strip()))
            heading_path = " > ".join(h[1] for h in heading_stack)
        else:
            heading_path = " > ".join(h[1] for h in heading_stack) if heading_stack else None

        chunks.append(
            ContentChunk(
                content=text.strip(),
                relative_path=file_path.name,
                chunk_index=len(chunks),
                heading_path=heading_path,
                paragraph_index=p_idx,
            )
        )

    return ParsedDocument(
        relative_path=file_path.name,
        format="docx",
        chunks=chunks,
    )


# ---------------------------------------------------------------------------
# XLSX parser (optional: openpyxl)
# ---------------------------------------------------------------------------


def _column_letter(n: int) -> str:
    """Convert 0-based column index to Excel column letter(s)."""
    result = ""
    while n >= 0:
        result = chr(65 + (n % 26)) + result
        n = n // 26 - 1
    return result


def parse_xlsx(file_path: Path) -> ParsedDocument:
    """Parse an XLSX file into per-sheet, per-row chunks.

    Requires ``openpyxl``.  The first non-empty row of each sheet is
    treated as the header.  Subsequent rows become chunks tagged with
    ``sheet_name``, ``header_columns``, and ``cell_range``.

    Returns:
        ParsedDocument with format ``"xlsx"``.
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError(
            "XLSX parsing requires openpyxl. Install with: pip install openpyxl"
        )

    wb = load_workbook(str(file_path), read_only=True, data_only=True)
    chunks: list[ContentChunk] = []

    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # Find first non-empty row as header
        header_idx = 0
        header_columns: list[str] = []
        for ri, row in enumerate(rows):
            values = [str(c) if c is not None else "" for c in row]
            if any(v.strip() for v in values):
                header_columns = values
                header_idx = ri
                break

        if not any(h.strip() for h in header_columns):
            header_columns = [f"Column{i}" for i in range(len(header_columns))]

        # Data rows
        for ri in range(header_idx + 1, len(rows)):
            values = [str(c) if c is not None else "" for c in rows[ri]]
            if not any(v.strip() for v in values):
                continue  # skip empty rows

            row_text_parts = []
            for col_i, val in enumerate(values):
                if val.strip():
                    hdr = header_columns[col_i] if col_i < len(header_columns) else f"Col{col_i}"
                    row_text_parts.append(f"{hdr}: {val}")

            row_text = " | ".join(row_text_parts)
            if not row_text.strip():
                continue

            start_col = _column_letter(0)
            end_col = _column_letter(max(len(values), len(header_columns)) - 1)
            start_row = ri + 1  # 1-indexed
            end_row = ri + 1
            cell_range = f"{start_col}{start_row}:{end_col}{end_row}"

            chunks.append(
                ContentChunk(
                    content=row_text,
                    relative_path=file_path.name,
                    chunk_index=len(chunks),
                    sheet_name=ws.title,
                    header_columns=header_columns,
                    cell_range=cell_range,
                )
            )

    wb.close()
    return ParsedDocument(
        relative_path=file_path.name,
        format="xlsx",
        chunks=chunks,
    )


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_PARSERS = {
    "md": parse_markdown,
    "txt": parse_txt,
    "pdf": parse_pdf,
    "docx": parse_docx,
    "xlsx": parse_xlsx,
}


def parse_file(file_path: Path, fmt: str) -> ParsedDocument:
    """Parse a single file into a ParsedDocument.

    Dispatches to the correct parser based on ``fmt``.

    Args:
        file_path: Absolute path to the file.
        fmt: Format string (``"md"``, ``"txt"``, ``"pdf"``, ``"docx"``, ``"xlsx"``).

    Returns:
        ParsedDocument with format-specific location chunks.

    Raises:
        ValueError: If ``fmt`` is not a supported format.
        FileNotFoundError: If the file does not exist.
        ImportError: If an optional dependency is missing (with install hint).
    """
    parser = _PARSERS.get(fmt)
    if parser is None:
        raise ValueError(f"Unsupported format: {fmt}")
    return parser(file_path)


# ---------------------------------------------------------------------------
# Full import pipeline
# ---------------------------------------------------------------------------


def import_project(project_id: str, *, base_dir: Path | None = None) -> ImportResult:
    """Run the full Phase A import pipeline for a project.

    1. Scan ``data/projects/<project_id>/source/``.
    2. Build SourceFile manifests (SHA-256, size, mtime).
    3. Parse each supported file; capture errors per-file.
    4. Return aggregated ImportResult.

    Args:
        project_id: Project identifier.
        base_dir: Root directory containing ``data/``. Defaults to ``Path.cwd()``.

    Returns:
        ImportResult with all file records and successfully parsed documents.
    """
    from app.rag.scanner import scan_source_dir

    base = base_dir or Path.cwd()

    try:
        scanned = scan_source_dir(project_id, base_dir=base)
    except FileNotFoundError:
        return ImportResult(
            project_id=project_id,
            total_files=0,
            success_count=0,
            failure_count=0,
            skipped_count=0,
        )

    files: list[SourceFile] = []
    parsed: list[ParsedDocument] = []
    success_count = 0
    failure_count = 0
    skipped_count = 0

    for full_path, rel_path, fmt in scanned:
        sf = build_source_file(full_path, rel_path, fmt)
        if fmt == "unsupported":
            sf.parse_status = "unsupported"
            skipped_count += 1
            files.append(sf)
            continue

        try:
            doc = parse_file(full_path, fmt)
            # Update relative_path to use full project-relative path
            doc.relative_path = rel_path
            for ch in doc.chunks:
                ch.relative_path = rel_path
            parsed.append(doc)
            sf.parse_status = "success"
            success_count += 1
        except Exception as exc:
            sf.parse_status = "failed"
            sf.error_message = f"{type(exc).__name__}: {exc}"
            failure_count += 1

        files.append(sf)

    return ImportResult(
        project_id=project_id,
        total_files=len(files),
        success_count=success_count,
        failure_count=failure_count,
        skipped_count=skipped_count,
        files=files,
        parsed=parsed,
    )
